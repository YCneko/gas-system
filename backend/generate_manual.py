"""将使用手册生成为 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re


doc = Document()

# ===== 页面设置 =====
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== 样式定义 =====
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = '微软雅黑'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        h_style.font.size = Pt(18)
        h_style.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        h_style.font.size = Pt(14)
        h_style.font.color.rgb = RGBColor(0, 80, 140)
    else:
        h_style.font.size = Pt(12)
        h_style.font.color.rgb = RGBColor(60, 60, 60)


def add_para(text, bold=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_code(text):
    """添加代码块 (灰色背景, 等宽字体)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    # 灰色背景
    shading_elm = p._element.get_or_add_pPr()
    shd = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0',
    })
    shading_elm.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p


def add_table(headers, rows, col_widths=None):
    """添加带格式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()  # spacing
    return table


def add_tip(text):
    """添加提示框"""
    p = doc.add_paragraph()
    shading_elm = p._element.get_or_add_pPr()
    shd = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'E8F5E9',
    })
    shading_elm.append(shd)
    run = p.add_run(f'💡 {text}')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_warning(text):
    """添加警告框"""
    p = doc.add_paragraph()
    shading_elm = p._element.get_or_add_pPr()
    shd = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'FFF3E0',
    })
    shading_elm.append(shd)
    run = p.add_run(f'⚠️ {text}')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_page_break():
    doc.add_page_break()


# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('Gas System', bold=True, size=36, color=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para('面向化工园区的多源废气智能治理系统', size=18, color=(80, 80, 80), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('— 使用手册 —', size=14, color=(120, 120, 120), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
add_para('零基础可用 | 10分钟上手', size=12, color=(100, 100, 100), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
add_para('版本 3.0 | 2026年6月', size=11, color=(150, 150, 150), align=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break()

# ============================================================
# 新手必读
# ============================================================
doc.add_heading('新手必读', level=1)

add_para('这份手册写给完全零基础的用户。你不需要懂编程知识，只要会：')
add_para('① 用浏览器打开网页')
add_para('② 在黑色命令窗口里复制粘贴文字')
add_para('③ 按回车键执行')
doc.add_paragraph()

add_para('这个系统是做什么的？', bold=True, size=14, color=(0, 51, 102))
add_para('一句话概括：帮化工厂的环保管理人员预测废气会不会超标。系统会自动做三件事：')
add_para('收集数据 — 把废气浓度、天气、设备运转情况自动汇总到一起')
add_para('预测未来 — 告诉你接下来 6 小时废气浓度会不会超标')
add_para('弹窗报警 — 如果真的快超标了，马上弹出警告提醒你')
doc.add_paragraph()

add_para('运行起来长什么样？', bold=True, size=14, color=(0, 51, 102))
add_para('最终你会看到这样一个网页：左侧有 4 个数据卡片（VOCs浓度、温度、湿度、风速），右侧上方是蓝黄相间的预测趋势图，右侧下方是预警列表。整体深色背景，类似数据大屏。')
doc.add_paragraph()

add_para('你需要准备什么？', bold=True, size=14, color=(0, 51, 102))
add_para('只需要一台能联网的 Windows 电脑。下面会一步步教你。')
add_page_break()

# ============================================================
# 方式一
# ============================================================
doc.add_heading('方式一：Docker 一键部署（推荐，最简单）', level=1)

add_tip('适用人群：完全零基础，想最快看到效果的用户。首次约 15 分钟（主要是下载时间），之后每次只需 30 秒。')

# 1.1
doc.add_heading('第 1 步：安装 Docker Desktop', level=2)
add_para('Docker 就像一个"箱子"，把整个系统打包进去，你不需要单独装 Python、MySQL 等任何东西。')
doc.add_paragraph()
add_para('操作步骤：', bold=True)
add_para('1. 打开浏览器，访问：https://www.docker.com/products/docker-desktop/')
add_para('2. 点击页面中间的蓝色大按钮 "Download Docker Desktop"')
add_para('3. 下载完成后，双击安装文件（文件名类似 Docker Desktop Installer.exe）')
add_para('4. 一路点"下一步" → "安装" → "完成"，全部用默认选项即可')
add_para('5. 安装完成后，桌面上会出现一个鲸鱼图标 🐳，双击打开它')
add_para('6. 第一次打开会要求注册账号，点 "Sign up" 用邮箱注册一个免费账号')
add_para('7. 注册完登录，等鲸鱼图标右下角变成绿色，就说明 Docker 启动好了')
doc.add_paragraph()
add_para('验证安装是否成功：', bold=True)
add_para('按键盘 Win+R，输入 cmd 回车，在黑色窗口里输入：')
add_code('docker --version')
add_para('如果显示类似 "Docker version 29.x.x"，说明安装成功。')

# 1.2
doc.add_heading('第 2 步：打开命令窗口', level=2)
add_para('1. 按键盘 Win + R，输入 cmd，点确定')
add_para('2. 一个黑色窗口会打开 — 这就是"命令提示符"，你只需要在里面复制粘贴命令即可')
add_para('3. 粘贴的快捷键是 Ctrl+V（或者右键 → 粘贴）')

# 1.3
doc.add_heading('第 3 步：进入项目目录', level=2)
add_para('把下面这行复制到黑色窗口里，按回车：')
add_code('cd /d "e:\\文件\\大二下作业\\Gas System"')

# 1.4
doc.add_heading('第 4 步：一键启动', level=2)
add_para('复制下面这条命令，粘贴到黑色窗口，按回车：')
add_code('docker-compose up -d --build')
add_warning('首次启动需要等待 5-15 分钟（取决于网速），因为它要下载约 500MB 的文件。屏幕上会有很多文字滚动，这是正常的。')
add_para('当看到类似下面的文字时，说明启动成功：')
add_code('Container gas_mysql    Started\nContainer gas_backend  Started\nContainer gas_frontend Started')

# 1.5
doc.add_heading('第 5 步：确认系统在运行', level=2)
add_para('在黑色窗口输入：')
add_code('docker ps')
add_para('你应该看到三个容器都在 "Up" 状态。只要都能看到，就说明一切正常！')

# 1.6
doc.add_heading('第 6 步：打开浏览器看效果', level=2)
add_para('1. 打开你电脑上的浏览器（Chrome、Edge、Firefox 随便哪个）')
add_para('2. 在地址栏输入：http://localhost')
add_para('3. 按回车')
doc.add_paragraph()
add_tip('看到深色背景的监控大屏，左边有 4 个数据卡片，右边有预测曲线图，说明部署成功！')

# 日常使用
doc.add_heading('以后怎么用？', level=2)
add_table(
    ['操作', '命令/方法'],
    [
        ['启动系统', '打开 Docker Desktop（鲸鱼图标变绿），然后重复第 2-4 步'],
        ['停止系统', '在命令窗口执行：docker-compose down'],
        ['完全退出', '打开 Docker Desktop → 右上角设置 → Quit Docker Desktop'],
    ]
)

add_page_break()

# ============================================================
# 方式二
# ============================================================
doc.add_heading('方式二：不用 Docker，手动启动（备选）', level=1)

add_tip('适用人群：电脑上已经有 Python 3.10+ 和 MySQL 8.0 的用户。首次约 20 分钟。')

doc.add_heading('前置准备：安装 Python', level=2)
add_para('如果你还没有安装 Python：')
add_para('1. 打开浏览器，访问：https://www.python.org/downloads/')
add_para('2. 页面会自动识别你的系统，点击黄色大按钮 "Download Python"')
add_para('3. 下载完成后双击安装')
add_warning('安装界面的最下面有个勾选框 "Add Python to PATH"，一定要勾上！否则后面命令无法识别 python 命令。')
add_para('4. 点 "Install Now"，等待安装完成')
doc.add_paragraph()
add_para('验证：', bold=True)
add_code('python --version')
add_para('显示 "Python 3.x.x" 就说明安装成功。')

doc.add_heading('前置准备：安装 MySQL', level=2)
add_para('1. 访问 https://dev.mysql.com/downloads/installer/')
add_para('2. 下载 mysql-installer-community（约 400MB）')
add_para('3. 双击安装，选择 "Developer Default"，一路下一步')
add_para('4. 在设置 root 密码时，设置为：123456')
add_para('5. 安装完成后，打开 MySQL 命令行，执行：')
add_code('CREATE DATABASE IF NOT EXISTS gas_system CHARACTER SET utf8mb4;')

doc.add_heading('第 1 步：进入后端目录', level=2)
add_code('cd /d "e:\\文件\\大二下作业\\Gas System\\backend"')

doc.add_heading('第 2 步：安装 Python 依赖包', level=2)
add_code('pip install -r requirements.txt')
add_warning('这步会比较慢（约 5-10 分钟），因为要下载 TensorFlow（约 350MB）。如果下载失败，换成国内镜像：')
add_code('pip install -r requirements.txt -i https://pypi.tuna.tsinghua.edu.cn/simple')

doc.add_heading('第 3 步：生成数据并训练模型', level=2)
add_para('依次执行下面三条命令（一条一条来，等上一条跑完再跑下一条）：')
doc.add_paragraph()
add_para('① 生成数据：', bold=True)
add_code('python generate_real_data.py')
add_para('看到 "Data import complete!" 说明成功。')
doc.add_paragraph()
add_para('② 数据分析：', bold=True)
add_code('python data_analysis.py')
add_para('看到 "特征矩阵 X 形状" 等字样说明成功，同时会生成一张热力图。')
doc.add_paragraph()
add_para('③ 训练模型：', bold=True)
add_code('python train_model.py')
add_para('约 1-2 分钟。屏幕上会显示训练进度，最后显示准确率 82.61%。')

doc.add_heading('第 4 步：启动后端', level=2)
add_code('python app.py')
add_para('此时命令窗口会显示：* Running on http://127.0.0.1:5000')
add_warning('这个窗口不要关！关了后端就停了。可以最小化到后台。')

doc.add_heading('第 5 步：新开命令窗口，启动前端', level=2)
add_para('再按 Win+R，输入 cmd，回车，打开第二个命令窗口。')
add_code('cd /d "e:\\文件\\大二下作业\\Gas System\\frontend"')
add_para('安装前端依赖（只需做一次）：')
add_code('npm install --registry=https://registry.npmmirror.com')
add_para('启动前端：')
add_code('npm run serve')
add_para('约 30 秒后看到 "Compiled successfully" 说明启动好了。')

doc.add_heading('第 6 步：打开浏览器', level=2)
add_para('在地址栏输入：http://localhost:8080')
add_tip('看到监控大屏说明成功！')

add_page_break()

# ============================================================
# 使用系统
# ============================================================
doc.add_heading('使用系统：浏览器操作指南', level=1)

doc.add_heading('看板页面说明', level=2)
add_table(
    ['区域', '位置', '内容', '怎么用'],
    [
        ['顶部标题栏', '最上面', '系统名称、时间、导出按钮', '点"导出报告"下载数据'],
        ['左侧卡片', '左边', 'VOCs浓度、温度、湿度、风速', '自动刷新，无需操作'],
        ['预测图表', '右上', '蓝色历史线 + 黄色预测线', '鼠标悬停看具体数值'],
        ['预警列表', '右下', '超标警告信息', '超标时自动出现，点某一行看详情'],
    ]
)

doc.add_heading('实时数据卡片（左侧四个框）', level=2)
add_para('自动每 5 秒刷新一次，显示最新的监测数值。如果数据一直是 0，说明数据库里还没有数据（参考"导入你自己的数据"章节）。')

doc.add_heading('预测趋势图（右上图表）', level=2)
add_table(
    ['线条颜色', '含义'],
    [
        ['蓝色实线', '过去 24 小时实际测到的浓度'],
        ['黄色虚线', 'AI 预测的未来 6 小时浓度'],
        ['红色虚线', '80 mg/m³ 预警线（超过这条线就会报警）'],
    ]
)
add_para('鼠标放到线上会弹出提示框，显示具体时间和浓度值。')

doc.add_heading('预警列表（右下表格）', level=2)
add_table(
    ['预警级别', '颜色', '含义', '触发条件'],
    [
        ['高', '🔴 红色', '严重，需立即处理', '超过 3 个预测值都超标'],
        ['中', '🟡 黄色', '注意，需关注', '2-3 个预测值超标'],
        ['低', '🔵 蓝色', '提醒，可关注', '1 个预测值超标'],
    ]
)
add_para('没有超标时，这个区域显示"暂无预警信息"。')

doc.add_heading('导出报告', level=2)
add_para('点击右上角的"导出报告"按钮，浏览器会自动下载一个 CSV 文件。这个文件用 Excel 就能直接打开，里面是历史排放数据。')

add_page_break()

# ============================================================
# 导入数据
# ============================================================
doc.add_heading('导入你自己的数据', level=1)
add_para('系统已经内置了 2160 条模拟数据。如果你想导入真实数据，有以下方式：')

doc.add_heading('方式 A：上传 CSV 或 Excel 文件（最简单）', level=2)
add_para('1. 把你的 CSV 或 Excel 文件准备好（格式参考下面的模板）')
add_para('2. 打开命令窗口，执行：')
add_code('curl -X POST http://localhost:5000/api/data/upload -F "file=@你的文件路径.csv" -F "data_type=emission"')
add_para('把"你的文件路径.csv"替换成你实际的文件路径。')
doc.add_paragraph()
add_para('data_type 参数有三个选项：', bold=True)
add_para('• emission — 废气排放数据')
add_para('• weather — 气象数据')
add_para('• equipment — 设备运行数据')

doc.add_heading('方式 B：用 API 实时发送数据', level=2)
add_para('适合有自动监测设备、想实时推送数据的场景：')
add_code('curl -X POST http://localhost:5000/api/data/realtime -H "Content-Type: application/json" -d "{\\"data_type\\":\\"emission\\",\\"records\\":[{\\"sensor_id\\":\\"S01\\",\\"timestamp\\":\\"2026-06-03 10:00:00\\",\\"voc_concentration\\":65.3}]}"')

doc.add_heading('CSV 文件模板', level=2)
add_para('你的 CSV 文件必须包含以下列（表头名称要一模一样）：')

add_para('废气排放数据 (emission)：', bold=True)
add_code('sensor_id,timestamp,voc_concentration,nox_concentration,so2_concentration\nS01,2026-06-03 10:00:00,120.5,45.6,10.2\nS01,2026-06-03 11:00:00,130.0,50.1,9.8')

add_para('气象数据 (weather)：', bold=True)
add_code('station_id,timestamp,temperature,humidity,wind_speed,wind_direction\nW01,2026-06-03 10:00:00,25.5,60.2,3.5,180')

add_para('设备数据 (equipment)：', bold=True)
add_code('equipment_id,timestamp,operating_load,production_phase,status\nE01,2026-06-03 10:00:00,72.5,正常,running')

add_tip('用 Excel 编辑好后，选择"文件 → 另存为 → CSV (逗号分隔)"即可。')

doc.add_heading('导入数据后需要重新训练吗？', level=2)
add_table(
    ['情况', '操作'],
    [
        ['只导入了少量数据（几十条）', '不需要，直接刷新浏览器页面就能看到新数据'],
        ['导入了大量新数据（几百条以上）', '建议重新训练。在 backend 目录下依次执行 python data_analysis.py 和 python train_model.py'],
        ['Docker 方式', '执行 docker-compose up -d --build backend'],
    ]
)

add_page_break()

# ============================================================
# 问题排查
# ============================================================
doc.add_heading('遇到问题怎么办？', level=1)

doc.add_heading('问题 1：浏览器打开显示"无法访问此网站"', level=2)
add_para('可能原因：服务没有启动。')
add_para('解决步骤：')
add_para('1. 打开命令窗口，输入 docker ps')
add_para('2. 如果看到空的表格或只有 1-2 行，说明服务没启动全')
add_para('3. 重新执行：cd /d "e:\\文件\\大二下作业\\Gas System" → docker-compose up -d')
add_para('4. 等 30 秒后刷新浏览器')

doc.add_heading('问题 2：Docker 启动报错 "port is already allocated"', level=2)
add_para('含义：端口被其他程序占用了。')
add_para('解决步骤：')
add_para('1. 用记事本打开 docker-compose.yml 文件')
add_para('2. 找到 "5000:5000"，改成 "5001:5000"')
add_para('3. 找到 "80:80"，改成 "8080:80"')
add_para('4. 保存文件后重新执行 docker-compose up -d')
add_para('5. 浏览器访问 http://localhost:8080 即可')

doc.add_heading('问题 3：页面上数据全部显示为 0', level=2)
add_para('含义：数据库里还没有数据。')
add_para('解决（Docker 方式）：')
add_code('docker exec -it gas_backend python generate_real_data.py')
add_para('解决（手动方式）：')
add_code('cd /d "e:\\文件\\大二下作业\\Gas System\\backend"\npython generate_real_data.py')

doc.add_heading('问题 4：pip install 下载失败', level=2)
add_para('提示 "Connection timeout" 或 "Could not find a version"。')
add_para('解决：使用国内镜像源：')
add_code('pip install -r requirements.txt -i https://pypi.tuna.tsinghua.edu.cn/simple')

doc.add_heading('问题 5：提示 "No module named xxx"', level=2)
add_para('含义：缺少某个 Python 包。')
add_para('解决：执行 pip install xxx（把 xxx 换成报错里提到的包名）。')

doc.add_heading('问题 6：docker-compose 命令不存在', level=2)
add_para('含义：Docker Desktop 没有正确安装或没有启动。')
add_para('解决：')
add_para('1. 确认桌面右下角有 Docker 鲸鱼图标')
add_para('2. 鲸鱼图标不是绿色 → 右键点它 → "Restart"')
add_para('3. 等图标变绿后再试')

add_page_break()

# ============================================================
# API 接口速查
# ============================================================
doc.add_heading('附录 A：API 接口速查表', level=1)

add_table(
    ['接口', '方法', '说明'],
    [
        ['/api/data/upload', 'POST', '上传 CSV/Excel 文件'],
        ['/api/data/realtime', 'POST', '接收 JSON 实时数据'],
        ['/api/realtime', 'GET', '获取最新监测数据'],
        ['/api/prediction', 'GET', '获取 24h历史 + 6h预测'],
        ['/api/vocs/prediction', 'GET', '预测 + 自动预警检查'],
        ['/api/alerts', 'GET', '查询预警记录列表'],
        ['/api/report/export', 'GET', '导出 CSV 报告'],
    ]
)

doc.add_heading('附录 B：数据库表字段速查', level=1)

add_para('emission_data（废气排放）', bold=True)
add_code('sensor_id | timestamp | voc_concentration | nox_concentration | so2_concentration')

add_para('weather_data（气象数据）', bold=True)
add_code('station_id | timestamp | temperature | humidity | wind_speed | wind_direction')

add_para('equipment_data（设备数据）', bold=True)
add_code('equipment_id | timestamp | operating_load | production_phase | status')

doc.add_heading('附录 C：名词解释', level=1)

add_table(
    ['名词', '解释'],
    [
        ['Docker', '一个打包工具，把整个系统装进"箱子"，不需要单独装各种依赖'],
        ['docker-compose', 'Docker 的遥控器，一条命令同时启动多个箱子'],
        ['容器 (Container)', 'Docker 箱子里的小隔间，每个隔间干一件事'],
        ['Flask', 'Python 写的网站后端框架，负责处理数据和业务逻辑'],
        ['API', '网站后台提供的数据接口，前端通过它获取数据'],
        ['Vue.js', '前端框架，负责页面的样子和交互效果'],
        ['ECharts', '画图表的工具，页面上那个好看的预测曲线图就是它画的'],
        ['LSTM', '一种 AI 算法，专门根据历史时序数据预测未来趋势'],
        ['TensorFlow', 'Google 的 AI 工具箱，LSTM 模型在它上面运行'],
        ['VOCs', '挥发性有机化合物，化工厂排放的主要废气污染物'],
        ['mg/m³', '毫克每立方米，衡量空气中污染物浓度的单位'],
        ['端口 (Port)', '电脑上的门牌号，不同服务用不同门牌号 (如 5000, 80)'],
        ['CSV', '最简单的表格文件格式，逗号分隔，Excel 能直接打开编辑'],
        ['JSON', '一种通用的数据交换格式，人和电脑都容易读懂'],
        ['curl', '命令行工具，用来在黑色窗口里测试网站接口'],
        ['localhost', '指你自己的电脑，127.0.0.1 也是同一个意思'],
        ['MySQL', '数据库软件，负责存储和管理所有数据'],
        ['Python', '编程语言，后端服务用它写的'],
        ['pip', 'Python 的软件包安装工具'],
        ['npm', 'Node.js 的软件包安装工具，前端用它安装依赖'],
    ]
)

add_page_break()

doc.add_heading('附录 D：项目文件结构', level=1)
add_code(
    'Gas-System/\n'
    '├── docker-compose.yml          Docker 编排配置\n'
    '├── 使用手册.md/docx            本使用手册\n'
    '├── backend/                    后端代码目录\n'
    '│   ├── app.py                  Flask 主程序 + 7 个 API 路由\n'
    '│   ├── models.py               数据库模型 (4 张表)\n'
    '│   ├── predictor.py            VOCSPredictor 推理类\n'
    '│   ├── alert.py                预警触发逻辑\n'
    '│   ├── data_cleaner.py         数据清洗工具\n'
    '│   ├── data_analysis.py        数据分析 + 特征工程\n'
    '│   ├── train_model.py          LSTM 模型训练 + 评估\n'
    '│   ├── generate_real_data.py   真实感数据生成器\n'
    '│   ├── test_integration.py     集成测试 + 性能评估\n'
    '│   ├── Dockerfile              后端 Docker 镜像\n'
    '│   └── requirements.txt        Python 依赖清单\n'
    '├── frontend/                   前端代码目录\n'
    '│   ├── Dockerfile              前端 Docker 镜像\n'
    '│   ├── nginx.conf              Nginx 配置\n'
    '│   ├── package.json            npm 依赖清单\n'
    '│   └── src/\n'
    '│       ├── views/DashboardView.vue   监控看板主页\n'
    '│       ├── components/               图表/卡片/预警组件\n'
    '│       └── utils/api.js             API 接口封装\n'
    '└── docker-compose.yml          整体编排配置\n'
)

# ============================================================
# 保存
# ============================================================
output_path = r'e:\文件\大二下作业\Gas System\使用手册.docx'
doc.save(output_path)
print(f'Word 文档已生成: {output_path}')
print('包含: 封面 + 新手必读 + Docker部署教程 + 手动部署教程 + 浏览器操作 + 数据导入 + 问题排查 + API速查 + 名词解释 + 项目结构')
